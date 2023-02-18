
import urllib.request
from rest_framework import viewsets, status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from .serializers import TodoSerializer
from .models import Todo, Files
from .serializers import FilesSerializer
from django.conf import settings
from urllib.parse import unquote
from django.http import JsonResponse
import docx
import os
import requests
from docx import Document
import time
import ssl

ssl._create_default_https_context = ssl._create_unverified_context

folder_name_origin = 'origin'
folder_name_split = 'split'
folder_name_wav = 'wav'
folder_name_docx = 'docx'
MEDIA_EXTERNAL = 'C:\\Users\\hungh\\OneDrive\\Desktop\\media'


# Create your views here.
@api_view(['GET'])
def translate(request, file_name, source, target):
    # create folder if not exist
    create_folder_if_not_exist(folder_name_split)  # to save file after splitting
    create_folder_if_not_exist(folder_name_origin)  # to save file origin
    texts = split_text(unquote(file_name))

    return JsonResponse({'text': unquote(texts), 'file_name': unquote(file_name)})


@api_view(['POST'])
def download_wav(request, file_name, folder_name):
    # create_folder_if_not_exist(folder_name_wav)  # to save file wav
    file_path = request.data['url'][14:]
    urllib.request.urlretrieve("http://studio.ploonet.com/" + file_path,
                               MEDIA_EXTERNAL + '/' + folder_name + '/' + file_name + ".wav")
    return Response('Success')


@api_view(['POST'])
def extract_url(request):
    url = request.data['url']
    text_obj = extract_text_from_url(url)
    return JsonResponse({'text': unquote(text_obj['text']), 'title': unquote(text_obj['title']),
                         'fileName': unquote(text_obj['file_name'])})


@api_view(['GET'])
def download(request, file_name):
    file_path = MEDIA_EXTERNAL + '/translated/' + unquote(file_name)
    file_pointer = open(file_path, 'r')
    response = Response(file_pointer,
                        content_type='application/vnd.openxmlformats-officedocument.presentationml.presentation')
    response['Content-Disposition'] = 'attachment; filename=' + unquote(file_name)
    return response


def extract_text_from_url(url):
    response = requests.get('http://13.21.34.201:8989/parser?url=' + url)
    if response.ok:
        text = response.json()['textContent']
        title = response.json()['title']
        image = response.json()['lead_image_url']
        return write_text_to_doc(text, title, url, image)
    return ''


def write_text_to_doc(texts, title, url, image):
    file_name = time.strftime("%Y%m%d-%H%M%S")
    create_folder_if_not_exist(file_name)
    # save image
    filename, file_extension = os.path.splitext(image)
    urllib.request.urlretrieve(image,
                               MEDIA_EXTERNAL + '/' + file_name + '/' + file_name + file_extension)
    document = Document()
    document.add_heading(title, level=1)
    document.add_heading('Source: ' + url, level=2)
    text_out = ''
    for text in texts.strip().split('.'):
        text_out += text.strip() + '.\n'
    document.add_paragraph(text_out)
    document.save(MEDIA_EXTERNAL + '/' + file_name + '/' + file_name + '.docx')
    return {'title': title, 'text': text_out, 'file_name': file_name}


def save_text_to_file(file_name, text):
    prefix_file_name = os.path.splitext(file_name)[0]
    # save file to temp folder in media folder
    file_path = MEDIA_EXTERNAL
    with open(file_path + '/temp/' + prefix_file_name + '.txt', "a") as myfile:
        myfile.write(text.strip() + "\n")


def create_folder_if_not_exist(folder_name):
    # crete folder if not exist in media folder
    print('======================== path = ' + MEDIA_EXTERNAL)
    if not os.path.exists(MEDIA_EXTERNAL + "/" + folder_name):
        os.mkdir(MEDIA_EXTERNAL + "/" + folder_name)


def split_text(file_name):
    file_path = MEDIA_EXTERNAL
    file_origin_path = file_path + '/' + folder_name_origin + '/' + file_name
    doc = docx.Document(file_origin_path)

    filename, file_extension = os.path.splitext(file_name)
    file_split_path = file_path + '/' + folder_name_split + '/' + filename + '.txt'
    texts = ''
    # create file if doesn't exist
    open(file_split_path, 'w+', encoding='utf-8')
    with open(file_split_path, "a", encoding='utf-8') as myfile:
        for p in doc.paragraphs:
            if len(p.text) != 0:
                texts += p.text
                myfile.write(p.text.strip())
    return texts

class TodoView(viewsets.ModelViewSet):
    serializer_class = TodoSerializer
    queryset = Todo.objects.all()


class FilesViewSet(viewsets.ModelViewSet):
    queryset = Files.objects.all()
    serializer_class = FilesSerializer
